import itertools
import sys
import argparse
import gzip
import re
from collections import defaultdict
from docx import Document
from docx.shared import Pt
from docx.shared import RGBColor
import xlsxwriter

parser = argparse.ArgumentParser(description='Summarize coverage for the painel of genes in analysis.')
parser.add_argument('-inputBed', '--inputBed', help='Output from mosdepht with the coverage of the genes with interest', required=True)
args = parser.parse_args()

## this search is optimized to find any hit, with the gene ORF4 they find MORF4L for example, the perfect match is done further because it is faster
def get_coverage(values, leng):
	dic={1:1, 2:5, 3:10, 4:15, 5:20, 6:30, 7:50, 8:100, 9:200} #intervals of mosdepht, this need to be change to be the same as the input file or     manually checked
	min_c=0
	max_c=0
	mean_c="{0:.2f}".format(0.00)
	bp_noc=0
	bp_noc15x="-"
	cov_per="{0:.2f}".format(0.00)

	if values[0]=="0":
		bp_noc=leng
		bp_noc15x=leng
		return leng, min_c, max_c, mean_c, cov_per, bp_noc, bp_noc15x
	else:
		laux=[x for x in values if x!= "0"]
		i=1
		a=0
		m=0
		max_c=dic.get(len(laux))
		for j in laux:
			m_aux=int(j)*int(dic.get(i))
			m=m_aux+m
			if int(j) >= int(a):
				min_c=dic.get(i)
			a=j
			i=i+1
		bp_noc=(leng-1-int(laux[0]))
		if len(laux)>=3:
			bp_noc15x=(leng-1-int(laux[2]))
		else:
			bp_noc15x=leng
		cov_per="{0:.2f}".format(((int(laux[0])+1)/float(leng))*100)
		mean_c="{0:.2f}".format((m/float(leng)))
		return leng, min_c, max_c, mean_c, cov_per, bp_noc, bp_noc15x

def getregionswithzerocov(linfo):
	lnotCov=[]
	diCov=dict() #only exons
	dinoCov=dict() #only exons
	for gene in linfo: #gene[15] is with no filter, gene[16] is with 1x,5x and 10x 
		if gene[16]!=0:
			if gene[4] not in lnotCov and gene[5] not in lnotCov:
				lnotCov.append((gene[0], gene[1], gene[2], gene[4], gene[16], gene[5]))
				if "E" in gene[5]:
					if gene[4] in dinoCov:
						dinoCov[gene[4]].append(gene[5])
					else:
						dinoCov[gene[4]]=[gene[5]]
		else:
			if "E" in gene[5]:
				if gene[4] in diCov and gene[5] not in diCov:
					diCov[gene[4]].append(gene[5])
				if gene[4] not in diCov:
					diCov[gene[4]]=[gene[5]]
	
	return lnotCov, diCov, dinoCov

def gatherInfoGenes(bedInput):
	print("Group the information ...")
	listInBed=[]
	firstLine=True
	lgenes=[]
	dic_reg={}

	with gzip.open(bedInput, 'rt') as fileIn:
		for line in fileIn:
			if firstLine and line.startswith("#"):
				line=fileIn.readline()
				firstLine=False
			else:
				firstLine=False
			auxline=line.rstrip().split("\t")
			chro=auxline[0]
			start=auxline[1]
			end=auxline[2]
			l=int(end)-int(start)+1
			gene=auxline[3].split(",")[0]
			if gene not in lgenes:
				lgenes.append(gene)
				dic_reg[gene]=(0,0)
			region=auxline[3].split(",")[1]
			if "E" in region and dic_reg[gene][0] == 0:
				e=int(region.rstrip().split("/")[1])
				dic_reg[gene]=(e,dic_reg[gene][1])
			if "I" in region and dic_reg[gene][1] == 0:
				i=int(region.rstrip().split("/")[1])
				dic_reg[gene]=(dic_reg[gene][0], i)	
			strand=auxline[3].split(",")[2]
			ensg=auxline[3].split(",")[3]
			enst=auxline[3].split(",")[4]
			rsmrna=auxline[3].split(",")[5]
			valcov=auxline[4:]
			leng, min_c, max_c, mean_c, cov_per, bp_noc, bp_noc15x=get_coverage(valcov, l)
			listInBed.append([chro, start, end, leng, gene, region, strand, ensg, enst, rsmrna, leng, min_c, max_c, mean_c, cov_per, bp_noc, bp_noc15x])
	lnotCov, dic_cov, dic_noCov=getregionswithzerocov(listInBed)
	sorted_lgenes=sorted(lgenes)
	sorted_lnotCov=sorted(lnotCov, key=lambda gene:gene[3])	
	sorted_listInBed=sorted(listInBed, key=lambda gene:gene[4])

	return sorted_lnotCov, dic_cov, dic_noCov, dic_reg, sorted_lgenes, sorted_listInBed

def constructword(listcov, dic_cov, dic_nocov, dic_regions, lgenes, outName):
	print("Number of genes: " + str(len(lgenes)))
	print("Genes: ")
	print(lgenes)
	print("Construct the word with the coverage summary ...")
	firstTable=True
	incomp=True
	latencao=[]

	outFileName=outName + ".docx"
	sampleName=outName.split(".")[0].replace("-coverage","")

	document=Document() #create the doc
	font=document.styles['Normal'].font
	document.add_heading("Resumo da cobertura do painel de genes analisado", 0)
	document.add_heading("Amostra:", level=2)
	document.add_paragraph(sampleName)
	document.add_heading("Painel de genes analisados:", level=2)
	if len(lgenes)<=10:
		for gene in lgenes:
			p=document.add_paragraph(gene)
			p.style='List Bullet'
	else:
		p=document.add_paragraph()
		for gene in lgenes:
			p.add_run("     ")
			p.add_run(gene)
	for gene in dic_nocov:
		for reg in dic_nocov[gene]:
			if gene in dic_cov:
				for value in dic_cov[gene]:
					if reg in value:
						aux=dic_nocov[gene]
						aux.remove(reg)
						dic_nocov[gene]=aux
			if gene not in dic_cov and gene not in latencao:
				latencao.append(gene)
	if latencao:
		run=document.add_paragraph().add_run('Atencao!')
		font=run.font
		font.bold=True
		font.color.rgb=RGBColor(214,24,24)
		font.size=Pt(13)
		document.add_heading('Genes com regioes que nao tiveram nenhuma base sequenciada: ', level=2)
		p=document.add_paragraph()
		for gene in latencao:
			p.add_run("     ")
			p.add_run(gene)
	for item in listcov:
		if firstTable:
			document.add_heading('Genes com regioes que nao foram totalmente sequenciadas: ', level=2)
			table_a=document.add_table(rows=1, cols=6)
			table_a.style = 'Light List Accent 1'
			hdr=table_a.rows[0].cells
			hdr[0].text="Chr"
			hdr[1].text="Start"
			hdr[2].text="End"
			hdr[3].text="Gene"
			hdr[4].text="BP not cov <15X"
			hdr[5].text="Region"
			firstTable=False
		rows=table_a.add_row().cells
		#if "I" in item[5]:
		#	rows[0].text=item[0]
		#	rows[1].text=item[1]
		#	rows[2].text=item[2]
		#	rows[3].text=item[3]
		#	rows[4].text=str(item[4])
		#	rows[5].text=item[5]
		if "E" in item[5]:
			cell=rows[1]
			cell.text=item[1]
			run=cell.paragraphs[0].runs[0]
		#	run.font.bold=True	
			cell=rows[2]	
			cell.text=item[2]
			run=cell.paragraphs[0].runs[0]
		#	run.font.bold=True
			cell=rows[3]
			cell.text=item[3]
			run=cell.paragraphs[0].runs[0]
		#	run.font.bold=True
			cell=rows[4]
			cell.text=str(item[4])
			run=cell.paragraphs[0].runs[0]
		#	run.font.bold=True
			cell=rows[5]
			cell.text=item[5]
			run=cell.paragraphs[0].runs[0]
		#	run.font.bold=True
			cell=rows[0]
			cell.text=item[0]
			run=cell.paragraphs[0].runs[0]
			run.font.bold=True	
	document.add_heading('Exoes totalmente sequenciadas: ', level=2)	
	for gene in dic_regions: #save all the introns[1] and exons[0], but use only the exons
		total=int(dic_regions[gene][0])#+int(dic_regions[gene][1])
		if gene in dic_nocov:
			nocov=len(dic_nocov[gene])
		if gene not in dic_nocov:
			nocov=0
		if gene in dic_cov:
			cov=len(dic_cov[gene])
		if gene not in dic_cov:
			cov=0
		total_cov=nocov+cov
		if cov==total_cov:
			seqed=str(total)+"/"+str(total)
			incomp=False
		else:
			if  total == total_cov:
				seqed=str(total-len(dic_nocov[gene]))+"/"+str(total)
			else:
				seqed=str(total-len(dic_nocov[gene]))+"/"+str(total)	
		if incomp:
			w=gene+": "+str(seqed)
			p=document.add_paragraph()
			p.add_run(w).bold=True
			p.style="List Bullet"
		else:
			p=document.add_paragraph(gene)
			p.add_run(": ")
			p.add_run(seqed)
			p.style='List Bullet'

		incomp=True
	document.save(outFileName)				

def constructexcel(l_bedIn, outFileName):
	print("Construct the excel ...")

	workbook=xlsxwriter.Workbook(outFileName + ".xlsx")
	worksheet = workbook.add_worksheet("Coverage")
	bold=workbook.add_format({'bold': True})
	worksheet.write('A1', 'Chromosome', bold)
	worksheet.write('B1', 'Start', bold)
	worksheet.write('C1', 'End', bold)
	worksheet.write('D1', 'Gene', bold)
	worksheet.write('E1', 'Length', bold)
	worksheet.write('F1', 'Min_Coverage', bold)
        worksheet.write('G1', 'Max_Coverage', bold)
        worksheet.write('H1', 'Mean_Coverage', bold)
        worksheet.write('I1', '%_Covered', bold)
	worksheet.write('J1', 'BP_notCovered', bold)
	worksheet.write('K1', 'BP_notCovered_<15X', bold)
	worksheet.write('L1', 'Region(+/- 20 bp)', bold) # this values need to be changed if the number of bp is changed before this step
	worksheet.write('M1', 'Strand', bold)
#	worksheet.write('N1', 'ENSGene', bold)
	worksheet.write('O1', 'ENSTranscript', bold)
	worksheet.write('P1', 'RefSeq_mRNA', bold)
	#start from the first cell below the header
	row =1
	col=0
	for i in l_bedIn:
		worksheet.write(row, col, i[0])
		worksheet.write(row, col + 1, i[1])
		worksheet.write(row, col + 2, i[2])
		worksheet.write(row, col + 3, i[4])
		worksheet.write(row, col + 4, i[3])
		worksheet.write(row, col + 5, i[11])
		worksheet.write(row, col + 6, i[12])	
		worksheet.write(row, col + 7, i[13])
		worksheet.write(row, col + 8, i[14])
		worksheet.write(row, col + 9, i[15])
		worksheet.write(row, col + 10, i[16])
		worksheet.write(row, col + 11, i[5])
		worksheet.write(row, col + 12, i[6])
#		worksheet.write(row, col + 13, i[7]) # ENSG removed from excel
		worksheet.write(row, col + 13, i[8])
		worksheet.write(row, col + 14, i[9])
		row+=1
	workbook.close()

if __name__=="__main__":

	bedFile=args.inputBed
	outname=("".join(bedFile.split('.')[0])+".summarize")
	
	t_bednotCov, dic_cov, dic_nocov, dic_regions, lgenes, listInBed=gatherInfoGenes(bedFile)
	constructexcel(listInBed, outname)
	constructword(t_bednotCov, dic_cov, dic_nocov, dic_regions, lgenes, outname)
